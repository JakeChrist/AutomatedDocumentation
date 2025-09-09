from pathlib import Path
from typing import Iterable
import importlib
import textwrap
import logging
import sys
import time
import json
import re

import pytest
from bs4 import BeautifulSoup

import explaincode
from explaincode import main
import manual_utils
from cache import ResponseCache


def _create_fixture(tmp_path: Path) -> None:
    nested = tmp_path / "subdir" / "nested"
    nested.mkdir(parents=True)
    (nested / "page.html").write_text(
        "<html><body><h1>Overview</h1></body></html>", encoding="utf-8"
    )
    content = textwrap.dedent(
        """
        # Overview
        Demo project

        # Purpose & Problem Solving
        Solves a problem

        # How to Run
        Usage: run it

        # Inputs
        Input data

        # Outputs
        Output data

        # System Requirements
        None

        # Examples
        Example usage
        """
    ).strip()
    (tmp_path / "README.md").write_text(content, encoding="utf-8")
    (tmp_path / "sample.json").write_text("{\"input\": \"data\"}", encoding="utf-8")


def _mock_llm_client() -> object:
    class Dummy:
        def summarize(self, text: str, prompt_type: str, system_prompt: str = "") -> str:  # pragma: no cover - simple stub
            return textwrap.dedent(
                """
                Overview: Demo project
                Purpose & Problem Solving: Solves a problem
                How to Run: Execute it
                Inputs: Input data
                Outputs: Output data
                System Requirements: None
                Examples: Example usage
                """
            ).strip()

    return Dummy()


def test_extract_text_markdown_preserves_headings_and_code(tmp_path: Path) -> None:
    content = """
# Title

Text

```python
print('hi')
```
"""
    md = tmp_path / "readme.md"
    md.write_text(content, encoding="utf-8")
    text = explaincode.extract_text(md)
    assert "# Title" in text
    assert "```python" in text
    assert "print('hi')" in text


def test_extract_text_html_preserves_headings_and_code(tmp_path: Path) -> None:
    html = (
        "<html><body><h1>Main</h1><p>Intro</p>"
        "<h2>Section</h2><pre><code>print('hi')</code></pre></body></html>"
    )
    page = tmp_path / "page.html"
    page.write_text(html, encoding="utf-8")
    text = explaincode.extract_text(page)
    assert "# Main" in text
    assert "## Section" in text
    assert "```" in text and "print('hi')" in text


def test_extract_text_docx_preserves_headings(tmp_path: Path) -> None:
    try:
        from docx import Document
    except Exception:  # pragma: no cover - dependency missing
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Text")
    path = tmp_path / "doc.docx"
    doc.save(path)
    text = explaincode.extract_text(path)
    assert text.splitlines()[0] == "# Title"
    assert "Text" in text


def test_render_html_renders_markdown_headings_and_code() -> None:
    sections = {
        "Intro": "# Title\n\n```python\nprint('hi')\n```",
    }
    html = explaincode.render_html(sections, "Manual")
    soup = BeautifulSoup(html, "html.parser")
    assert soup.find("h1", string="Title") is not None
    code = soup.find("pre").find("code")
    assert code is not None and "print('hi')" in code.text


def test_render_html_includes_toc_and_sources_block() -> None:
    sections = {"Overview": "", "How to Run": ""}
    evidence = {
        "Overview": {"evidence": [{"snippet": "info", "file": "readme.md"}]},
        "How to Run": {"evidence": [{"snippet": "run", "file": "run.py"}]},
    }
    html = explaincode.render_html(sections, "Manual", evidence)
    soup = BeautifulSoup(html, "html.parser")
    nav = soup.find("nav")
    assert nav is not None and nav.find("a", href="#overview") is not None
    sources_blocks = soup.find_all("div", class_="sources")
    assert any("readme.md" in block.text for block in sources_blocks)
    assert any("run.py" in block.text for block in sources_blocks)


def test_html_summary_creation(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _create_fixture(tmp_path)
    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    main(["--path", str(tmp_path)])
    manual = tmp_path / "user_manual.html"
    evidence = tmp_path / "user_manual_evidence.json"
    assert manual.exists()
    assert evidence.exists()
    data = json.loads(evidence.read_text(encoding="utf-8"))
    assert "Overview" in data
    assert data["Overview"]["evidence"]
    html_text = manual.read_text(encoding="utf-8")
    assert "No information provided." not in html_text
    soup = BeautifulSoup(html_text, "html.parser")
    nav = soup.find("nav")
    assert nav is not None and nav.find("a", href="#overview") is not None


def test_pdf_summary_creation(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    if importlib.util.find_spec("reportlab") is None:
        pytest.skip("reportlab not installed")
    _create_fixture(tmp_path)
    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    main(["--path", str(tmp_path), "--output-format", "pdf"])
    assert (tmp_path / "user_manual.pdf").exists()
    assert (tmp_path / "user_manual_evidence.json").exists()


def test_graceful_missing_docx(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    _create_fixture(tmp_path)
    try:
        from docx import Document
    except Exception:  # pragma: no cover - dependency missing
        Document = None  # type: ignore
    if Document:
        doc = Document()
        doc.add_paragraph("hi")
        doc.save(tmp_path / "guide.docx")
    monkeypatch.setattr(explaincode, "Document", None)
    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    main(["--path", str(tmp_path)])
    assert (tmp_path / "user_manual.html").exists()
    assert (tmp_path / "user_manual_evidence.json").exists()


def test_custom_output_directory(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _create_fixture(tmp_path)
    out_dir = tmp_path / "dist"
    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    main(["--path", str(tmp_path), "--output", str(out_dir)])
    assert (out_dir / "user_manual.html").exists()
    assert (out_dir / "user_manual_evidence.json").exists()


def test_collect_docs_filters(tmp_path: Path) -> None:
    (tmp_path / "docs").mkdir()
    (tmp_path / "docs" / "keep.md").write_text("hi", encoding="utf-8")
    (tmp_path / "docs" / "skip.txt").write_text("no", encoding="utf-8")
    (tmp_path / "README.md").write_text("readme", encoding="utf-8")
    (tmp_path / "extra.md").write_text("extra", encoding="utf-8")
    files = explaincode.collect_docs(tmp_path)
    names = {f.name for f in files}
    assert "keep.md" in names and "skip.txt" not in names


def test_map_evidence_overview_priority_and_filters() -> None:
    docs = {
        Path("README.md"): "# Overview\nreadme info",
        Path("docs/guide.md"): "# Overview\ndocs info",
        Path("src/other.md"): "# Overview\nother info",
        Path("tests/ignore.md"): "# Overview\ntest info",
        Path("examples/ignore.md"): "# Overview\nexample info",
        Path("fixtures/ignore.md"): "# Overview\nfixture info",
    }
    section_map, _ = explaincode.map_evidence_to_sections(docs)
    sources = [p.as_posix() for p, _ in section_map["Overview"]]
    assert "tests/ignore.md" not in sources
    assert "examples/ignore.md" not in sources
    assert "fixtures/ignore.md" not in sources
    assert set(sources[:2]) == {"README.md", "docs/guide.md"}
    assert sources[2] == "src/other.md"


def test_map_evidence_snippet_limits() -> None:
    long_content = "# Inputs\n" + "\n".join(
        f"line {i}" for i in range(1, explaincode.MAX_SNIPPET_LINES + 10)
    )
    docs = {
        Path("src/long.md"): long_content,
        Path("tests/short.md"): long_content,
    }
    section_map, _ = explaincode.map_evidence_to_sections(docs)
    snippets = {p.as_posix(): snip for p, snip in section_map["Inputs"]}
    assert (
        f"line {explaincode.MAX_SNIPPET_LINES + 1}" not in snippets["src/long.md"]
    )
    assert f"line {explaincode.MAX_SNIPPET_LINES}" in snippets["src/long.md"]
    assert snippets["tests/short.md"] == "# Inputs"


def test_detect_placeholders() -> None:
    text = "Overview: [[NEEDS_OVERVIEW]]\nInputs: data\nOutputs: [[NEEDS_OUTPUTS]]"
    missing = explaincode.detect_placeholders(text)
    assert set(missing) == {"Overview", "Outputs"}


def test_parse_manual_infers_missing_sections() -> None:
    class Stub:
        def summarize(self, prompt: str, prompt_type: str, system_prompt: str = "") -> str:
            return "guessed"

    client = Stub()
    parsed = explaincode.parse_manual("Overview: hi", client=client)
    assert parsed["Overview"] == "hi"
    assert parsed["Inputs"].endswith("(inferred)")
    assert "No information provided." not in parsed["Inputs"]


def test_validate_manual_references_flags_missing(tmp_path: Path) -> None:
    (tmp_path / "exists.py").write_text("pass", encoding="utf-8")
    sections = {"Overview": "See exists.py and missing.py for details"}
    evidence: dict[str, dict[str, object]] = {}
    explaincode.validate_manual_references(sections, tmp_path, evidence)
    assert "exists.py" in sections["Overview"]
    assert "missing.py [missing]" in sections["Overview"]
    assert evidence["Overview"]["missing_references"] == ["missing.py"]


def test_infer_sections_infers_entries() -> None:
    sections = explaincode.infer_sections("Some context")
    assert sections["Overview"] == "Some context"
    for key in explaincode.REQUIRED_SECTIONS:
        if key != "Overview":
            assert sections[key].endswith("(inferred)")
    assert "No information provided." not in "".join(sections.values())


def test_infer_sections_no_context_defaults() -> None:
    sections = explaincode.infer_sections("")
    assert all(v == "No information provided." for v in sections.values())


def test_extract_snippets_skips_large_file(
    tmp_path: Path, caplog: pytest.LogCaptureFixture
) -> None:
    big_file = tmp_path / "big.py"
    big_file.write_bytes(b"a" * 210_000)
    caplog.set_level(logging.INFO)
    snippets = explaincode.extract_snippets(
        [big_file], max_files=1, time_budget=5, max_bytes=200_000
    )
    assert big_file not in snippets
    log = caplog.text
    assert "file size" in log and "exceeds limit" in log


def test_scan_code_skips_non_source_dirs(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    src = tmp_path / "src"
    src.mkdir()
    (src / "main.py").write_text('"""run code"""', encoding="utf-8")

    tests_dir = tmp_path / "tests"
    tests_dir.mkdir()
    (tests_dir / "test_main.py").write_text('"""run tests"""', encoding="utf-8")

    examples_dir = tmp_path / "examples"
    examples_dir.mkdir()
    (examples_dir / "example.py").write_text('"""run example"""', encoding="utf-8")

    monkeypatch.setattr(explaincode, "collect_docs", lambda base: [])
    result = explaincode.scan_code(
        tmp_path, ["How to Run"], max_files=5, time_budget=5, max_bytes_per_file=1000
    )

    how_to_run = result.get("How to Run", {})
    assert set(how_to_run.keys()) == {"src/main.py"}


def test_scan_code_categorizes_snippets(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    paths = [tmp_path / "a.py", tmp_path / "b.py", tmp_path / "c.py"]
    for p in paths:
        p.write_text("", encoding="utf-8")

    monkeypatch.setattr(explaincode, "collect_docs", lambda base: [])
    monkeypatch.setattr(explaincode, "rank_code_files", lambda base, patterns: paths)

    def fake_extract(
        files: Iterable[Path], *, max_files: int, time_budget: int, max_bytes: int
    ) -> dict[Path, str]:
        return {
            paths[0]: "read input from user",
            paths[1]: "write output",
            paths[2]: "run the tool",
        }

    monkeypatch.setattr(explaincode, "extract_snippets", fake_extract)

    result = explaincode.scan_code(
        tmp_path, ["Inputs", "Outputs", "How to Run"], max_files=3, time_budget=5, max_bytes_per_file=1000
    )
    assert result["Inputs"] == {"a.py": "read input from user"}
    assert result["Outputs"] == {"b.py": "write output"}
    assert result["How to Run"] == {"c.py": "run the tool"}


def test_rank_code_files_supports_cpp_h_java(tmp_path: Path) -> None:
    files = [
        tmp_path / "main.py",
        tmp_path / "lib.cpp",
        tmp_path / "lib.h",
        tmp_path / "Main.java",
        tmp_path / "ignore.txt",
    ]
    for f in files:
        f.write_text("", encoding="utf-8")
    ranked = explaincode.rank_code_files(tmp_path, [])
    names = {p.name for p in ranked}
    assert {"main.py", "lib.cpp", "lib.h", "Main.java"} <= names
    assert "ignore.txt" not in names


def test_llm_fill_placeholders_per_section_logging(
    tmp_path: Path, caplog: pytest.LogCaptureFixture
) -> None:
    manual = "Inputs: [[NEEDS_INPUTS]]\nOutputs: [[NEEDS_OUTPUTS]]"
    evidence = {
        "Inputs": {"in.py": "input data"},
        "Outputs": {"out.py": "output data"},
    }

    class Dummy:
        def __init__(self) -> None:
            self.calls: list[str] = []

        def summarize(
            self, text: str, prompt_type: str, system_prompt: str = ""
        ) -> str:
            self.calls.append(text)
            manual_match = re.search(r"Manual:\n(.*?)\n\nSection:", text, re.DOTALL)
            section_match = re.search(r"Section: (.*?)\n", text)
            manual_text = manual_match.group(1)
            section = section_match.group(1)
            token = explaincode.SECTION_PLACEHOLDERS[section]
            return manual_text.replace(token, f"filled {section}")

    client = Dummy()
    cache = ResponseCache(str(tmp_path / "cache.json"))
    caplog.set_level(logging.INFO)
    result = explaincode.llm_fill_placeholders(manual, evidence, client, cache)
    assert "filled Inputs" in result and "filled Outputs" in result
    assert len(client.calls) == 2
    log = caplog.text
    assert "Filled Inputs using code from: in.py" in log
    assert "Filled Outputs using code from: out.py" in log




def test_full_docs_no_code_scan(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, caplog: pytest.LogCaptureFixture
) -> None:
    _create_fixture(tmp_path)
    tracker = {"rank": 0, "extract": 0}

    def fake_rank(root: Path, patterns: list[str]) -> list[Path]:
        tracker["rank"] += 1
        return [tmp_path / "a.py"]

    def fake_extract(
        files: Iterable[Path], *, max_files: int, time_budget: int, max_bytes: int
    ) -> dict[Path, str]:
        tracker["extract"] += 1
        return {}

    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    monkeypatch.setattr(explaincode, "rank_code_files", fake_rank)
    monkeypatch.setattr(explaincode, "extract_snippets", fake_extract)

    caplog.set_level(logging.INFO)
    main(["--path", str(tmp_path), "--scan-code-if-needed"])
    assert tracker["rank"] == 0
    assert tracker["extract"] == 0
    log = caplog.text
    assert "DOC PASS" in log
    assert "Code scan skipped" in log


def test_missing_run_triggers_code_fallback_with_limits(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, caplog: pytest.LogCaptureFixture
) -> None:
    (tmp_path / "README.md").write_text("Only overview", encoding="utf-8")

    class Dummy:
        def summarize(self, text: str, prompt_type: str, system_prompt: str = "") -> str:
            if "How to Run" in system_prompt:
                return "[[NEEDS_RUN_INSTRUCTIONS]]"
            if "enhancing a user manual" in system_prompt:
                manual_match = re.search(r"Manual:\n(.*?)\n\nSection:", text, re.DOTALL)
                manual_text = manual_match.group(1)
                return manual_text.replace("[[NEEDS_RUN_INSTRUCTIONS]]", "use it")
            return "x"

    paths = [tmp_path / f"f{i}.py" for i in range(3)]
    tracker: dict[str, object] = {"rank": 0, "extract": 0}

    def fake_rank(root: Path, patterns: list[str]) -> list[Path]:
        tracker["rank"] += 1
        return paths

    def fake_extract(
        files: Iterable[Path], *, max_files: int, time_budget: int, max_bytes: int
    ) -> dict[Path, str]:
        tracker["extract"] += 1
        lst = list(files)
        tracker["kwargs"] = {"max_files": max_files, "time_budget": time_budget}
        tracker["scanned"] = lst[:max_files]
        return {p: "run code" for p in lst[:max_files]}

    monkeypatch.setattr(explaincode, "LLMClient", lambda: Dummy())
    monkeypatch.setattr(explaincode, "rank_code_files", fake_rank)
    monkeypatch.setattr(explaincode, "extract_snippets", fake_extract)

    caplog.set_level(logging.INFO)
    main(
        [
            "--path",
            str(tmp_path),
            "--scan-code-if-needed",
            "--max-code-files",
            "1",
            "--code-time-budget-seconds",
            "5",
        ]
    )
    html = (tmp_path / "user_manual.html").read_text(encoding="utf-8")
    assert "NEEDS_RUN_INSTRUCTIONS" not in html
    assert tracker["rank"] == 1
    assert tracker["extract"] == 1
    assert tracker["kwargs"] == {"max_files": 1, "time_budget": 5}
    assert len(tracker["scanned"]) == 1
    log = caplog.text
    assert "Pass 1 missing sections" in log
    assert "How to Run" in log
    assert "Filled How to Run using code from: f0.py" in log


def test_no_code_flag_skips_code_fallback(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, caplog: pytest.LogCaptureFixture
) -> None:
    (tmp_path / "README.md").write_text("Only overview", encoding="utf-8")

    class Dummy:
        def summarize(
            self, text: str, prompt_type: str, system_prompt: str = ""
        ) -> str:
            return "Overview: x\\nHow to Run: [[NEEDS_RUN_INSTRUCTIONS]]"

    tracker = {"rank": 0, "extract": 0}

    def fake_rank(root: Path, patterns: list[str]) -> list[Path]:
        tracker["rank"] += 1
        return []

    def fake_extract(
        files: Iterable[Path], *, max_files: int, time_budget: int, max_bytes: int
    ) -> dict[Path, str]:
        tracker["extract"] += 1
        return {}

    monkeypatch.setattr(explaincode, "LLMClient", lambda: Dummy())
    monkeypatch.setattr(explaincode, "rank_code_files", fake_rank)
    monkeypatch.setattr(explaincode, "extract_snippets", fake_extract)

    caplog.set_level(logging.INFO)
    main(["--path", str(tmp_path), "--scan-code-if-needed", "--no-code"])
    assert tracker["rank"] == 0
    assert tracker["extract"] == 0
    log = caplog.text
    assert "Code scan skipped: --no-code specified" in log
    assert "How to Run" in log


def test_force_code_flag_triggers_code_fallback(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, caplog: pytest.LogCaptureFixture
) -> None:
    _create_fixture(tmp_path)
    tracker = {"rank": 0, "extract": 0}

    def fake_rank(root: Path, patterns: list[str]) -> list[Path]:
        tracker["rank"] += 1
        return [tmp_path / "script.py"]

    def fake_extract(
        files: Iterable[Path], *, max_files: int, time_budget: int, max_bytes: int
    ) -> dict[Path, str]:
        tracker["extract"] += 1
        return {next(iter(files)): "code"}

    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    monkeypatch.setattr(explaincode, "rank_code_files", fake_rank)
    monkeypatch.setattr(explaincode, "extract_snippets", fake_extract)

    caplog.set_level(logging.INFO)
    main(["--path", str(tmp_path), "--force-code"])
    assert tracker["rank"] == 1
    assert tracker["extract"] == 1
    log = caplog.text
    assert "Code scan triggered: --force-code enabled" in log
    assert "Pass 1 complete: no sections missing" in log


def test_custom_title_and_filename(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _create_fixture(tmp_path)
    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    main(["--path", str(tmp_path), "--title", "Fancy Guide"])
    out_file = tmp_path / "fancy_guide.html"
    evidence_file = tmp_path / "fancy_guide_evidence.json"
    assert out_file.exists()
    assert evidence_file.exists()
    assert "<h1>Fancy Guide</h1>" in out_file.read_text(encoding="utf-8")


def test_insert_into_index(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _create_fixture(tmp_path)
    out_dir = tmp_path / "out"
    out_dir.mkdir()
    index = out_dir / "index.html"
    index.write_text("<html><body><ul></ul></body></html>", encoding="utf-8")
    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    main(["--path", str(tmp_path), "--output", str(out_dir), "--insert-into-index"])
    html = index.read_text(encoding="utf-8")
    assert '<a href="user_manual.html">User Manual</a>' in html


def test_docs_index_default_and_injection(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    _create_fixture(tmp_path)
    docs_dir = tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "index.html").write_text("<html><body><nav></nav></body></html>", encoding="utf-8")
    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    main(["--path", str(tmp_path), "--insert-into-index"])
    manual = docs_dir / "user_manual.html"
    evidence = docs_dir / "user_manual_evidence.json"
    assert manual.exists()
    assert evidence.exists()
    soup = BeautifulSoup((docs_dir / "index.html").read_text(encoding="utf-8"), "html.parser")
    nav = soup.find("nav")
    assert nav is not None
    first = nav.find("a")
    assert first and first["href"] == "user_manual.html"


def test_insert_into_root_index(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _create_fixture(tmp_path)
    (tmp_path / "index.html").write_text("<html><body><nav></nav></body></html>", encoding="utf-8")
    monkeypatch.setattr(explaincode, "LLMClient", _mock_llm_client)
    main(["--path", str(tmp_path), "--insert-into-index"])
    manual = tmp_path / "docs" / "user_manual.html"
    evidence = tmp_path / "docs" / "user_manual_evidence.json"
    assert manual.exists()
    assert evidence.exists()
    soup = BeautifulSoup((tmp_path / "index.html").read_text(encoding="utf-8"), "html.parser")
    nav = soup.find("nav")
    assert nav is not None
    first = nav.find("a")
    assert first and first["href"] == "docs/user_manual.html"


def test_chunking_triggers_multiple_calls_and_logs(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    paragraph1 = "aaa " * 2000
    paragraph2 = "bbb " * 2000
    text = "\n\n".join([paragraph1, paragraph2])

    class Dummy:
        def __init__(self) -> None:
            self.calls: list[dict] = []

        def summarize(self, text: str, prompt_type: str, system_prompt: str = "") -> str:
            self.calls.append(
                {"text": text, "prompt_type": prompt_type, "system_prompt": system_prompt}
            )
            if system_prompt == explaincode.MERGE_SYSTEM_PROMPT:
                return "final"
            return text.split()[0]

    client = Dummy()
    cache = ResponseCache(str(tmp_path / "cache.json"))
    logging.basicConfig(stream=sys.stdout, level=logging.DEBUG, force=True)
    result = manual_utils._summarize_manual(
        client, cache, text, chunking="auto", source="src"
    )

    assert result == "final"
    assert len(client.calls) == 3
    chunk_calls = client.calls[:-1]
    merge_call = client.calls[-1]
    assert all(c["system_prompt"] == explaincode.CHUNK_SYSTEM_PROMPT for c in chunk_calls)
    assert merge_call["system_prompt"] == explaincode.MERGE_SYSTEM_PROMPT
    assert merge_call["text"] == "aaa\n\nbbb"

    out = capsys.readouterr().out
    assert "Chunk 1/2 from src" in out
    assert "Merged LLM response length" in out


def test_chunk_edit_hook_applied(tmp_path: Path) -> None:
    paragraph1 = "aaa " * 2000
    paragraph2 = "bbb " * 2000
    text = "\n\n".join([paragraph1, paragraph2])

    class Dummy:
        def __init__(self) -> None:
            self.calls: list[dict] = []

        def summarize(self, text: str, prompt_type: str, system_prompt: str = "") -> str:
            self.calls.append(
                {"text": text, "prompt_type": prompt_type, "system_prompt": system_prompt}
            )
            if system_prompt == explaincode.MERGE_SYSTEM_PROMPT:
                return "final"
            return text.split()[0]

    def hook(chunks: list[str]) -> list[str]:
        return [c.upper() for c in chunks]

    client = Dummy()
    cache = ResponseCache(str(tmp_path / "cache.json"))
    result = manual_utils._summarize_manual(
        client, cache, text, chunking="auto", source="src", post_chunk_hook=hook
    )

    assert result == "final"
    assert client.calls[-1]["text"] == "AAA\n\nBBB"


def test_parallel_chunk_summarization(tmp_path: Path) -> None:
    paragraph = "word " * 2000
    text = "\n\n".join([paragraph, paragraph])
    delay = 0.2

    class SlowClient:
        def summarize(self, text: str, prompt_type: str, system_prompt: str = "") -> str:
            if system_prompt == explaincode.CHUNK_SYSTEM_PROMPT:
                time.sleep(delay)
            return "ok"

    client = SlowClient()
    cache = ResponseCache(str(tmp_path / "cache.json"))
    start = time.perf_counter()
    manual_utils._summarize_manual(client, cache, text, chunking="auto", source="src")
    duration = time.perf_counter() - start
    assert duration < delay * 1.5


def test_hierarchical_merge_logged(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    big = "word " * 5000
    text = "\n\n".join([big, big])

    class Dummy:
        def __init__(self) -> None:
            self.calls: list[dict] = []

        def summarize(self, text: str, prompt_type: str, system_prompt: str = "") -> str:
            self.calls.append(
                {"text": text, "prompt_type": prompt_type, "system_prompt": system_prompt}
            )
            if len(self.calls) <= 2:
                return big
            return "short"

    client = Dummy()
    cache = ResponseCache(str(tmp_path / "cache.json"))
    logging.basicConfig(stream=sys.stdout, level=logging.DEBUG, force=True)
    result = manual_utils._summarize_manual(
        client, cache, text, chunking="auto", source="src"
    )

    assert result == "short"
    assert len(client.calls) > 3
    out = capsys.readouterr().out
    assert "Hierarchical merge pass" in out


def test_cached_chunks_reused(tmp_path: Path) -> None:
    paragraph = "word " * 2000
    text = "\n\n".join([paragraph, paragraph])

    class Dummy:
        def __init__(self) -> None:
            self.calls: list[dict] = []

        def summarize(self, text: str, prompt_type: str, system_prompt: str = "") -> str:
            self.calls.append(
                {"text": text, "prompt_type": prompt_type, "system_prompt": system_prompt}
            )
            return f"resp{len(self.calls)}"

    cache_file = tmp_path / "cache.json"
    cache = ResponseCache(str(cache_file))
    client1 = Dummy()
    result1 = manual_utils._summarize_manual(
        client1, cache, text, chunking="auto", source="src"
    )
    assert result1 == "resp3"
    assert len(client1.calls) == 3

    client2 = Dummy()
    cache2 = ResponseCache(str(cache_file))
    result2 = manual_utils._summarize_manual(
        client2, cache2, text, chunking="auto", source="src"
    )
    assert result2 == "resp3"
    assert len(client2.calls) == 0


def test_chunking_none_no_llm_calls(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    big_text = "data " * 5000
    (tmp_path / "README.md").write_text("# Overview\n" + big_text, encoding="utf-8")

    class Dummy:
        def __init__(self) -> None:
            self.calls: list[dict] = []

        def summarize(self, text: str, prompt_type: str, system_prompt: str = "") -> str:
            self.calls.append(
                {"text": text, "prompt_type": prompt_type, "system_prompt": system_prompt}
            )
            return "done"

    dummy = Dummy()
    monkeypatch.setattr(explaincode, "LLMClient", lambda: dummy)
    main(["--path", str(tmp_path), "--chunking", "none"])

    assert len(dummy.calls) == 5
    call = dummy.calls[0]
    assert "Overview" in call["system_prompt"]


def test_llm_generate_manual_sanitizes_and_caches(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    docs = {Path("readme.md"): "# Overview\nExample"}
    cache = ResponseCache(str(tmp_path / "cache.json"))
    client = explaincode.LLMClient("http://fake")

    def fake_summarize(prompt: str, prompt_type: str, system_prompt: str = "") -> str:
        return (
            "You are a documentation engine. Summarize the following.\n"
            "It prints output."
        )

    monkeypatch.setattr(client, "summarize", fake_summarize)
    manual, _, _ = explaincode.llm_generate_manual(docs, client, cache, chunking="none")
    lower = manual.lower()
    assert "documentation engine" not in lower
    assert "summarize the following" not in lower

    section_map, _ = explaincode.map_evidence_to_sections(docs)
    entries = section_map["Overview"]
    context = "\n\n".join(snippet for _, snippet in entries)
    prompt = (
        "Write the 'Overview' section of a user manual using the "
        "following documentation snippets.\n\n" + context
    )
    key = ResponseCache.make_key("section:Overview", prompt)
    cached = cache.get(key)
    assert cached is not None
    assert "documentation engine" not in cached.lower()
